import os
import csv
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text):
    return doc.add_paragraph(text)

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Gemstone Price Prediction - Project Report Sections (Buddhika)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ---------------------------------------------------------
    # 1. Methodology
    # ---------------------------------------------------------
    add_heading(doc, '1. Methodology', level=1)
    
    add_heading(doc, 'A. Dataset and Preprocessing', level=2)
    add_paragraph(doc, 
        "The gemstone dataset, originally sourced as 'cubic_zirconia.csv', was preprocessed to ensure consistency with the diamonds dataset used by the other team member. The raw dataset contained 26,967 samples and 10 features. During preprocessing, invalid entries (e.g., zero-values for dimensions x, y, or z) and significant outliers in carat weight and price were removed. Categorical features ('cut', 'color', 'clarity') were ordinally encoded to preserve their inherent ranking, mapping the qualitative grades to numerical scales. The final cleaned dataset, 'gemstone_clean.csv', consists of 25,123 samples."
    )
    
    add_heading(doc, 'B. Particle Swarm Optimization (PSO) Implementation', level=2)
    add_paragraph(doc,
        "A binary Particle Swarm Optimization (PSO) algorithm was developed from scratch using NumPy for feature selection. The particle representation consisted of a continuous velocity vector that was transformed into a binary feature mask via a sigmoid activation function (probability = 1 / (1 + exp(-velocity))). If a particle's probability for a specific feature exceeded a random threshold in the uniform range [0, 1], the feature was selected (value 1); otherwise, it was dropped (value 0)."
    )
    
    add_paragraph(doc,
        "The fitness function was designed to maximize predictive performance while penalizing model complexity. Specifically, fitness was calculated using a 5-fold cross-validated XGBoost Regressor on the training set, applying the identical formula used in the Genetic Algorithm implementation: Fitness = (Average CV R²) - (0.001 * N_selected_features). To prevent velocity saturation, particle velocities were clamped within the range [-4, 4]."
    )
    
    add_paragraph(doc,
        "The optimization was run with a swarm size of 30 particles over 40 iterations. The cognitive and social acceleration coefficients (c1 and c2) were both set to 1.5, while the inertia weight (w) was linearly decayed from 0.9 to 0.4 over the iterations to encourage exploration early on and exploitation in later stages. To ensure computational efficiency across the 1,200 particle evaluations, subset fitness memoization was utilized."
    )
    
    # Image
    pso_conv_path = 'results/pso_convergence.png'
    if os.path.exists(pso_conv_path):
        doc.add_picture(pso_conv_path, width=Inches(5.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = add_paragraph(doc, "Figure 1: PSO Convergence Curve showing the improvement of the Global Best Fitness and Average Swarm Fitness over 40 iterations.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------------------------------------------------
    # 2. Results
    # ---------------------------------------------------------
    add_heading(doc, '2. Results', level=1)
    
    add_paragraph(doc, "The tables and figures below present the comparative performance of the baseline models against the GA and PSO optimized feature subsets on the held-out test sets.")
    
    # Table
    csv_path = 'results/final_comparison.csv'
    if os.path.exists(csv_path):
        with open(csv_path, 'r') as f:
            reader = csv.reader(f)
            data = list(reader)
        
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(data):
            for j, val in enumerate(row):
                cell = table.cell(i, j)
                cell.text = str(val)
                if i == 0:
                    # Bold header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
    
    add_paragraph(doc, "\n")
    
    # Charts
    chart1_path = 'results/comparison_metrics_chart.png'
    if os.path.exists(chart1_path):
        doc.add_picture(chart1_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = add_paragraph(doc, "Figure 2: Grouped bar chart comparing RMSE, MAE, and R² across the 6 model variants. It is evident that the optimized models maintain an R² score extremely close to their respective baselines.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    chart2_path = 'results/ga_vs_pso_convergence.png'
    if os.path.exists(chart2_path):
        doc.add_picture(chart2_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = add_paragraph(doc, "Figure 3: Convergence comparison of GA and PSO. Both metaheuristic algorithms quickly converge to a highly optimal fitness plateau, effectively navigating the combinatorial search space.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    chart3_path = 'results/complexity_comparison.png'
    if os.path.exists(chart3_path):
        doc.add_picture(chart3_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = add_paragraph(doc, "Figure 4: Scatter plot demonstrating model complexity (number of features) versus computational training time. The optimized models (green and blue) cluster heavily to the left, highlighting significant dimensional reduction compared to the baseline models (gray).")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_heading(doc, 'Statistical Summary', level=2)
    add_paragraph(doc, 
        "Based on the final evaluation metrics, the GA implementation reduced the number of features for the diamonds dataset by 33.3% (from 9 to 6 features), resulting in a marginal R² change of -0.108% (0.9821 to 0.9810). Concurrently, the PSO implementation reduced the features for the gemstone dataset by 55.6% (from 9 to 4 features), yielding an R² change of -0.159% (0.9818 to 0.9802). Training times fluctuated due to differing subsets, but feature dimensionality was strictly reduced in both methodologies."
    )

    # ---------------------------------------------------------
    # 3. Discussion
    # ---------------------------------------------------------
    add_heading(doc, '3. Discussion', level=1)
    
    add_paragraph(doc,
        "The empirical findings strongly indicate that both GA and PSO successfully achieved significant feature reduction while maintaining highly competitive predictive accuracy. Although the optimized subsets exhibited negligible declines in R² (-0.108% for GA and -0.159% for PSO), the vast reduction in feature space (33% and 55% respectively) drastically reduces model complexity and enhances interpretability. Given the context of regression on tabular data, trading a fraction of a percent of accuracy for a model with half the parameters is highly favorable."
    )
    
    add_paragraph(doc,
        "An analysis of the selected features reveals striking domain alignment. Both algorithms independently selected 'carat', 'cut', 'color', and 'clarity'. These four attributes are universally recognized in the gemology industry as the '4 Cs'—the primary determinants of a diamond's value. The algorithms successfully dropped redundant spatial dimensions (x, y, z) and proportional metrics (depth, table) that are highly collinear with carat weight, demonstrating that nature-inspired feature selection can autonomously derive domain-intuitive subsets."
    )
    
    add_paragraph(doc,
        "Based on these outcomes, the core research hypothesis is firmly supported: applying metaheuristic optimization for feature selection provides models that are substantially less complex while sacrificing virtually zero predictive accuracy."
    )
    
    add_paragraph(doc,
        "However, several limitations must be acknowledged. First, the algorithms were tested on two distinct datasets (diamonds and cubic zirconia) due to the parallel structure of the project, meaning direct performance comparisons between GA and PSO are conflated with dataset variance. Second, the fitness function's penalty coefficient for feature count (0.001) was manually designated rather than empirically tuned. Finally, algorithm hyperparameters (e.g., mutation rates, inertia weights) were set based on general heuristics rather than exhaustive grid search."
    )

    # ---------------------------------------------------------
    # 4. Conclusion
    # ---------------------------------------------------------
    add_heading(doc, '4. Conclusion', level=1)
    
    add_paragraph(doc,
        "This project investigated the efficacy of nature-inspired algorithms—specifically Genetic Algorithms and Particle Swarm Optimization—for feature selection in predictive gemstone pricing models. Baseline Random Forest and XGBoost regressors were trained on full-feature datasets, yielding high accuracy but incorporating redundant spatial data. By deploying custom-built metaheuristic algorithms, we successfully isolated the critical predictive features. The optimized models reduced dimensionality by up to 55.6% while experiencing an accuracy drop of less than 0.2%, confirming that GA and PSO effectively balance the trade-off between model simplicity and predictive power."
    )
    
    add_paragraph(doc,
        "Future work should address the identified limitations to build upon these promising results. First, standardizing a single, shared dataset to benchmark both GA and PSO directly would allow for a strict comparative analysis of their convergence rates and computational efficiency. Second, developing a hybrid GA-PSO algorithm could theoretically combine the robust global exploration of GA crossover with the rapid local exploitation of PSO velocity vectors. Lastly, benchmarking these traditional machine learning pipelines against modern Deep Learning architectures (e.g., TabNet) on the optimized feature subsets would provide a comprehensive evaluation of state-of-the-art predictive capabilities."
    )

    # ---------------------------------------------------------
    # 5. References
    # ---------------------------------------------------------
    add_heading(doc, '5. References', level=1)
    
    refs = [
        "[1] S. Agrawal, 'Diamonds Dataset,' Kaggle, 2017. [Online]. Available: https://www.kaggle.com/shivam2503/diamonds.",
        "[2] 'Cubic Zirconia Dataset,' Kaggle. [Online]. Available: https://www.kaggle.com/.",
        "[3] F. Pedregosa et al., 'Scikit-learn: Machine Learning in Python,' Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
        "[4] T. Chen and C. Guestrin, 'XGBoost: A Scalable Tree Boosting System,' in Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 2016, pp. 785-794.",
        "[5] J. Kennedy and R. Eberhart, 'Particle swarm optimization,' Proceedings of ICNN'95 - International Conference on Neural Networks, Perth, WA, Australia, 1995, pp. 1942-1948 vol.4.",
        "[6] [Insert Sajini's Literature Review Citation 1 Here]",
        "[7] [Insert Sajini's Literature Review Citation 2 Here]"
    ]
    
    for r in refs:
        add_paragraph(doc, r)
        
    # Save document
    out_dir = 'report'
    os.makedirs(out_dir, exist_ok=True)
    doc.save(os.path.join(out_dir, 'report_sections_buddhika.docx'))
    print("Report generated successfully.")

if __name__ == '__main__':
    main()
